from __future__ import annotations

import html
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, PageBreak


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
TIMES = Path("C:/Windows/Fonts/times.ttf")
TIMES_BOLD = Path("C:/Windows/Fonts/timesbd.ttf")


def read_markdown(path: Path) -> list[str]:
    return path.read_text(encoding="utf-8").splitlines()


def normalize_inline(text: str) -> str:
    escaped = html.escape(text)
    escaped = escaped.replace("`", "")
    return escaped


def add_docx_paragraph(document: Document, text: str, style: str | None = None):
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def markdown_to_docx(source: Path, target: Path):
    document = Document()
    section = document.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(13)
    styles["Heading 1"].font.name = "Times New Roman"
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 2"].font.name = "Times New Roman"
    styles["Heading 2"].font.size = Pt(16)
    styles["Heading 3"].font.name = "Times New Roman"
    styles["Heading 3"].font.size = Pt(14)

    in_code = False
    for raw_line in read_markdown(source):
        line = raw_line.rstrip()
        if line == "<!-- pagebreak -->":
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            continue
        if line.startswith("```"):
            in_code = not in_code
            continue
        if not line:
            continue
        if line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            add_docx_paragraph(document, line[2:], style="List Bullet")
        elif line.startswith("> "):
            add_docx_paragraph(document, line[2:], style="Quote")
        elif in_code:
            add_docx_paragraph(document, line)
        else:
            add_docx_paragraph(document, line)

    document.save(target)


def register_pdf_fonts():
    if TIMES.exists():
      pdfmetrics.registerFont(TTFont("TimesNewRoman", str(TIMES)))
    if TIMES_BOLD.exists():
      pdfmetrics.registerFont(TTFont("TimesNewRoman-Bold", str(TIMES_BOLD)))


def markdown_to_pdf(source: Path, target: Path):
    register_pdf_fonts()
    font_name = "TimesNewRoman" if TIMES.exists() else "Helvetica"
    bold_font_name = "TimesNewRoman-Bold" if TIMES_BOLD.exists() else "Helvetica-Bold"

    styles = getSampleStyleSheet()
    normal = ParagraphStyle(
        "AuraNormal",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=13,
        leading=19.5,
        spaceAfter=8,
    )
    heading1 = ParagraphStyle(
        "AuraH1",
        parent=styles["Heading1"],
        fontName=bold_font_name,
        fontSize=18,
        leading=24,
        spaceBefore=10,
        spaceAfter=12,
    )
    heading2 = ParagraphStyle(
        "AuraH2",
        parent=styles["Heading2"],
        fontName=bold_font_name,
        fontSize=16,
        leading=22,
        spaceBefore=10,
        spaceAfter=10,
    )
    heading3 = ParagraphStyle(
        "AuraH3",
        parent=styles["Heading3"],
        fontName=bold_font_name,
        fontSize=14,
        leading=20,
        spaceBefore=8,
        spaceAfter=8,
    )
    bullet = ParagraphStyle(
        "AuraBullet",
        parent=normal,
        leftIndent=18,
        bulletIndent=6,
    )

    story = []
    in_code = False
    for raw_line in read_markdown(source):
        line = raw_line.rstrip()
        if line == "<!-- pagebreak -->":
            story.append(PageBreak())
            continue
        if line.startswith("```"):
            in_code = not in_code
            continue
        if not line:
            story.append(Spacer(1, 0.12 * cm))
            continue
        if line.startswith("# "):
            story.append(Paragraph(normalize_inline(line[2:]), heading1))
        elif line.startswith("## "):
            story.append(Paragraph(normalize_inline(line[3:]), heading2))
        elif line.startswith("### "):
            story.append(Paragraph(normalize_inline(line[4:]), heading3))
        elif line.startswith("- "):
            story.append(Paragraph(normalize_inline(line[2:]), bullet, bulletText="•"))
        elif line.startswith("> "):
            story.append(Paragraph(normalize_inline(line[2:]), normal))
        elif in_code:
            story.append(Paragraph(normalize_inline(line), normal))
        else:
            story.append(Paragraph(normalize_inline(line), normal))

    pdf = SimpleDocTemplate(
        str(target),
        pagesize=A4,
        rightMargin=2.2 * cm,
        leftMargin=2.2 * cm,
        topMargin=2.0 * cm,
        bottomMargin=2.0 * cm,
        title=source.stem,
    )
    pdf.build(story)


def main():
    documents = [
        ("final-report.md", "AuraCare-AI-Bao-Cao-Cuoi-Ky"),
        ("ai-prompts-appendix.md", "AuraCare-AI-Phu-Luc-Prompts"),
    ]
    for markdown_name, output_name in documents:
        source = DOCS / markdown_name
        markdown_to_docx(source, DOCS / f"{output_name}.docx")
        markdown_to_pdf(source, DOCS / f"{output_name}.pdf")
        print(f"generated {output_name}.docx/.pdf")


if __name__ == "__main__":
    main()
